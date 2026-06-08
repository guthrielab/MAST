#!/usr/bin/env python3
"""
M. tuberculosis tNGS resistance-prediction pipeline.

Renders the branded clinical report (tNGS_Report_Template.docx) with docxtpl.

CLI (matches the Nextflow process; trailing args are optional for stand-alone use):
    compare_mutations.py <variants.tsv> <id> <patient_dir> <reference.fasta> \
                         [mutations_csv] [lineage_csv] [template_docx] \
                         [patient_info_csv] [bam]
"""

import os
import re
import sys
from copy import deepcopy
from collections import defaultdict

import pandas as pd
import pysam
from docx.table import _Row
from docx.shared import RGBColor
from docxtpl import DocxTemplate, RichText

# --------------------------------------------------------------------------- #
#  INputs
# --------------------------------------------------------------------------- #
input_file         = sys.argv[1]
output_base_name   = sys.argv[2]
patient_dir        = sys.argv[3]
fasta_file         = sys.argv[4]
mutations_csv_path = sys.argv[5] if len(sys.argv) > 5 else '../../../MAST/Data/all_resistant_variants.csv'
lineage_csv_path   = sys.argv[6] if len(sys.argv) > 6 else '../../../MAST/Data/Lineage.csv'
template_path      = sys.argv[7] if len(sys.argv) > 7 else '../../../MAST/Data/tNGS_Report_Template.docx'
patient_info_path  = sys.argv[8] if len(sys.argv) > 8 else '../../../MAST/Data/patient_info.csv'
bam_arg            = sys.argv[9] if len(sys.argv) > 9 else None
regions_bed        = sys.argv[10] if len(sys.argv) > 10 else os.environ.get('TB_REGIONS_BED', 'regions.bed')
QC_MIN_DEPTH       = float(os.environ.get('TB_QC_DEPTH', '30'))

# --------------------------------------------------------------------------- #
#  Amino-acid abbrev conversion
# --------------------------------------------------------------------------- #
AA_MAP = {
    'Ala': 'A', 'Arg': 'R', 'Asn': 'N', 'Asp': 'D', 'Cys': 'C',
    'Gln': 'Q', 'Glu': 'E', 'Gly': 'G', 'His': 'H', 'Ile': 'I',
    'Leu': 'L', 'Lys': 'K', 'Met': 'M', 'Phe': 'F', 'Pro': 'P',
    'Ser': 'S', 'Thr': 'T', 'Trp': 'W', 'Tyr': 'Y', 'Val': 'V',
}


def format_mutation(mutation):
    """database format -> display format (e.g. rpoB_p.Ser450Tyr -> rpoB (S450Y))."""
    if '_LoF' in mutation:
        return f"{mutation.split('_')[0]} (LoF*)"
    if 'fs' in mutation and '_p.' in mutation:
        gene, aa_part = mutation.split('_p.')
        m = re.match(r'([A-Za-z]+)(\d+)(fs)', aa_part)
        if m:
            aa_full, pos, _ = m.groups()
            return f"{gene} ({AA_MAP.get(aa_full, aa_full[:1])}{pos}fs**)"
    if '_p.' in mutation:
        gene, aa_change = mutation.split('_p.')
        m = re.match(r'([A-Za-z]+)(\d+)([A-Za-z]+)', aa_change)
        if m:
            a_from, pos, a_to = m.groups()
            return f"{gene} ({AA_MAP.get(a_from, a_from[:1])}{pos}{AA_MAP.get(a_to, a_to[:1])})"
    return mutation


def split_gene_code(formatted):
    """'rpoB (S450Y)' -> ('rpoB', 'S450Y'); fallback ('', formatted)."""
    m = re.match(r'(\S+)\s*\((.+)\)', formatted)
    return (m.group(1), m.group(2)) if m else ('', formatted)


CONF_RANK = {'High': 3, 'Moderate': 2, 'Low': 1, 'N/A': 0, '': 0}
DEFAULT_CONF = 'Moderate'

LINEAGE_DESC = {
    '1': '(Indo-Oceanic lineage)', '2': '(East Asian lineage)',
    '2.2': '(East Asian/Beijing lineage)', '3': '(East African-Indian lineage)',
    '4': '(Euro-American lineage)', '4.9': '(Euro-American lineage/H37Rv-like)',
    '5': '(West African 1 lineage)', '6': '(West African 2 lineage)',
}

# --------------------------------------------------------------------------- #
#  Load inputs
# --------------------------------------------------------------------------- #
df           = pd.read_csv(input_file, sep='\t')
df_mutations = pd.read_csv(mutations_csv_path)
df_lineage   = pd.read_csv(lineage_csv_path)

df.columns           = df.columns.str.upper()
df_mutations.columns = df_mutations.columns.str.upper()
df_lineage.columns   = df_lineage.columns.str.upper()

HAS_CONF = 'CONFIDENCE' in df_mutations.columns

# --------------------------------------------------------------------------- #
#  Per-variant consensus -> confidence  (>=0.9 High, >0.8 Moderate)
# --------------------------------------------------------------------------- #
FREQ_CANDIDATES = ['CONSENSUS', 'VAF', 'ALT_FREQ', 'ALT_FREQUENCY',
                   'FREQ', 'FREQUENCY', 'ALLELE_FREQUENCY']
FREQ_COL = os.environ.get('TB_FREQ_COL', '').upper() or \
           next((c for c in FREQ_CANDIDATES if c in df.columns), None)
HAS_INFO = 'INFO' in df.columns
print(f"Consensus source: "
      f"{FREQ_COL or ('INFO (AO/DP)' if HAS_INFO else 'BAM' if bam_arg else 'none')}")

# use BAM for consensus look-ups
cons_bam = None
if bam_arg and os.path.exists(bam_arg):
    try:
        if not os.path.exists(bam_arg + '.bai'):
            pysam.index(bam_arg)
        cons_bam = pysam.AlignmentFile(bam_arg, "rb")
    except Exception as e:
        print(f"Could not open BAM for consensus: {e}")


def to_fraction(value):
    try:
        x = float(value)
    except (TypeError, ValueError):
        return None
    return x / 100.0 if x > 1 else x


def consensus_from_info(info):
    """Observed allele fraction = AO/DP parsed from a freebayes INFO string."""
    d = {}
    for kv in str(info).split(';'):
        if '=' in kv:
            k, v = kv.split('=', 1)
            d[k] = v
    try:
        if 'AO' in d and 'DP' in d:
            ao = float(str(d['AO']).split(',')[0])   # first ALT (NUMALT=1 here)
            dp = float(d['DP'])
            if dp > 0:
                return ao / dp
    except (ValueError, ZeroDivisionError):
        pass
    return None


def bam_consensus(pos1, ref, alt):
    """Fraction of reads supporting a single-base ALT at a 1-based position."""
    if cons_bam is None or len(ref) != 1 or len(alt) != 1:
        return None
    try:
        cov = cons_bam.count_coverage(
            cons_bam.references[0], pos1 - 1, pos1, quality_threshold=0)
        counts = {'A': cov[0][0], 'C': cov[1][0], 'G': cov[2][0], 'T': cov[3][0]}
        total = sum(counts.values())
        return counts.get(alt.upper(), 0) / total if total else None
    except Exception:
        return None


def get_consensus(tsv_row):
    """Best available observed consensus (0-1) for a matched variant row."""
    if FREQ_COL:
        c = to_fraction(tsv_row[FREQ_COL])
        if c is not None:
            return c
    if HAS_INFO:
        c = consensus_from_info(tsv_row['INFO'])
        if c is not None:
            return c
    return bam_consensus(int(tsv_row['POS']), str(tsv_row['REF']), str(tsv_row['ALT']))


def confidence_label(cons, db_conf):
    """Map a numeric consensus to High/Moderate/Low; else fall back to DB/default."""
    if cons is not None:
        if cons >= 0.9:
            return 'High'
        if cons > 0.8:
            return 'Moderate'
        return 'Low'
    return db_conf or DEFAULT_CONF


def variants_match(tsv_pos, tsv_ref, tsv_alt, csv_pos, csv_ref, csv_alt):
    if tsv_pos == csv_pos and tsv_ref == csv_ref and tsv_alt == csv_alt:
        return True
    if len(tsv_ref) > len(tsv_alt) and abs(tsv_pos - csv_pos) <= 1 and len(csv_ref) > len(csv_alt):
        td, cd = tsv_ref[len(tsv_alt):], csv_ref[len(csv_alt):]
        if td in cd or cd in td:
            return True
    if len(tsv_alt) > len(tsv_ref) and abs(tsv_pos - csv_pos) <= 1 and len(csv_alt) > len(csv_ref):
        ti, ci = tsv_alt[len(tsv_ref):], csv_alt[len(csv_ref):]
        if ti in ci or ci in ti:
            return True
    if abs(tsv_pos - csv_pos) <= 1:
        if (tsv_ref in csv_ref or csv_ref in tsv_ref) and (tsv_alt in csv_alt or csv_alt in tsv_alt):
            return True
    return False


# --------------------------------------------------------------------------- #
#  Match resistance variants
# --------------------------------------------------------------------------- #
hits = defaultdict(dict)   # drug -> {variant: {'cons': float|None, 'db': str|None}}

for _, tsv_row in df.iterrows():
    for _, csv_row in df_mutations.iterrows():
        if variants_match(tsv_row['POS'], tsv_row['REF'], tsv_row['ALT'],
                           csv_row['POSITION'], csv_row['REFERENCE_NUCLEOTIDE'],
                           csv_row['ALTERNATIVE_NUCLEOTIDE']):
            drug = csv_row['DRUG']
            variant = csv_row['VARIANT']

            cons = get_consensus(tsv_row)
            db_conf = str(csv_row['CONFIDENCE']) if HAS_CONF else None

            rec = hits[drug].get(variant, {'cons': None, 'db': db_conf})
            if cons is not None and (rec['cons'] is None or cons > rec['cons']):
                rec['cons'] = cons          # keep the strongest observation
            rec['db'] = rec['db'] or db_conf
            hits[drug][variant] = rec

            print(f"Matched {variant} ({drug}) at {tsv_row['POS']} "
                  f"{tsv_row['REF']}>{tsv_row['ALT']}  consensus="
                  f"{cons if cons is not None else 'NA'}")

if cons_bam is not None:
    cons_bam.close()

# list of variants per drug (preserves first-seen order)
resistances = {drug: list(vmap.keys()) for drug, vmap in hits.items()}

# --------------------------------------------------------------------------- #
#  Lineage detection
# --------------------------------------------------------------------------- #
merged_lin = pd.merge(
    df, df_lineage,
    left_on=['POS', 'REF', 'ALT'],
    right_on=['POSITION', 'REFERENCE_NUCLEOTIDE', 'ALTERNATIVE_NUCLEOTIDE'],
    how='inner',
)

lineage_value, lineage_detected = 'Unknown', False
if not merged_lin.empty:
    lineage_value, lineage_detected = merged_lin['LIN'].iloc[0], True
    print(f"Lineage detected from mutation list: {lineage_value}")

if not lineage_detected:
    lineage_reference_positions = {420008: ('4.9', 'A')}

    # Prefer the BAM passed in by Nextflow; otherwise fall back to a search.
    bam_file = None
    if bam_arg and os.path.exists(bam_arg):
        bam_file = bam_arg
        print(f"Using staged BAM file: {bam_file}")
    else:
        print("No BAM argument provided/found; searching work dirs ...")
        search_dirs = ['.', patient_dir]
        current_dir = os.getcwd()
        for i in range(1, 4):
            p = os.path.abspath(os.path.join(current_dir, *['..'] * i))
            if os.path.exists(p):
                search_dirs.append(p)
        search_dirs = list({d for d in search_dirs if os.path.exists(d)})
        for search_dir in search_dirs:
            try:
                for root, _, files in os.walk(search_dir):
                    for f in files:
                        if f.endswith('.bam') and output_base_name in f and 'sorted' in f.lower():
                            cand = os.path.join(root, f)
                            try:
                                with pysam.AlignmentFile(cand, "rb") as tb:
                                    if tb.nreferences > 0:
                                        bam_file = cand
                                        break
                            except Exception:
                                continue
                    if bam_file:
                        break
            except (PermissionError, OSError):
                continue
            if bam_file:
                break

    if bam_file and os.path.exists(bam_file):
        print(f"Lineage reference check using: {bam_file}")
        try:
            if not os.path.exists(bam_file + '.bai'):
                pysam.index(bam_file)
            bam = pysam.AlignmentFile(bam_file, "rb")
            for pos, (lineage, expected_base) in lineage_reference_positions.items():
                cov = bam.count_coverage(contig=bam.references[0], start=pos - 1, stop=pos)
                counts = {'A': cov[0][0], 'C': cov[1][0], 'G': cov[2][0], 'T': cov[3][0]}
                total = sum(counts.values())
                print(f"Position {pos}: coverage={total}, bases={counts}")
                if total >= 10:
                    consensus = max(counts, key=counts.get)
                    if consensus == expected_base and counts[consensus] / total >= 0.9:
                        lr = df_lineage[df_lineage['POSITION'] == pos]
                        lineage_value = lr['LIN'].iloc[0] if not lr.empty else lineage
                        lineage_detected = True
                        print(f"Detected lineage {lineage_value} from ref position {pos}")
                        break
            bam.close()
        except Exception as e:
            print(f"Error processing BAM file {bam_file}: {e}")

    if not lineage_detected:
        print("No lineage detected from mutations or reference positions")

# --------------------------------------------------------------------------- #
#  Resolve output dir + patient record
# --------------------------------------------------------------------------- #
os.makedirs(patient_dir, exist_ok=True)

df_pat = pd.read_csv(patient_info_path)
row = df_pat[df_pat['Barcode'] == output_base_name]
if row.empty:
    print(f"No record found with Barcode: {output_base_name}", file=sys.stderr)
    sys.exit(1)
patient_info = row.to_dict(orient='records')[0]

# --------------------------------------------------------------------------- #
#  Build context (scalars + colour-coded RichText)
# --------------------------------------------------------------------------- #
DRUGS = ['Ethambutol', 'Ethionamide', 'Pyrazinamide', 'Isoniazid', 'Rifampicin',
         'Streptomycin', 'Ciprofloxacin', 'Ofloxacin', 'Moxifloxacin', 'Amikacin',
         'Kanamycin', 'Capreomycin', 'Bedaquiline', 'Linezolid']
FIRST_LINE = {'Ethambutol', 'Pyrazinamide', 'Isoniazid', 'Rifampicin'}
FLUOROQUINOLONES = {'Moxifloxacin', 'Ofloxacin'}
NEW_SECONDLINE = {'Linezolid', 'Bedaquiline'}


def yes_rt():
    rt = RichText(); rt.add('Yes', bold=True, color='C00000'); return rt


def no_rt():
    rt = RichText(); rt.add('No', color='2E7D32'); return rt


context = dict(patient_info)
context['Lineage'] = lineage_value
context['Lineage_desc'] = LINEAGE_DESC.get(str(lineage_value), '')

findings = []
tsv_data = {'Sample': output_base_name, 'Lineage': lineage_value}

for drug in DRUGS:
    variants = resistances.get(drug, [])
    if variants:
        vmap = hits[drug]                       # {variant: {'cons','db'}}

        # ---- pick ONE mutation for the Resistance Summary: highest confidence ----
        def rank_key(v):
            rec = vmap[v]
            has_cons = rec['cons'] is not None
            cons = rec['cons'] if has_cons else 0.0
            db = CONF_RANK.get(rec['db'] or '', 0)
            return (has_cons, cons, db)         # consensus-bearing first, then value

        best = max(variants, key=rank_key)
        best_label = confidence_label(vmap[best]['cons'], vmap[best]['db'])

        context[f'{drug}_detected'] = yes_rt()
        context[f'{drug}_gene'] = format_mutation(best)     # single mutation only
        context[f'{drug}_conf'] = best_label

        # ---- Detailed Findings keeps the FULL list of mutations ----
        formatted = [format_mutation(v) for v in variants]
        genes, codes = zip(*(split_gene_code(f) for f in formatted))
        if len(formatted) > 1:                              # multiple takes priority
            comment = 'Multiple mutations detected'
        elif any('LoF' in c for c in codes):
            comment = 'LoF mutation detected'
        else:
            comment = 'Mutation detected'
        findings.append({
            'drug': drug,
            'mut': '; '.join(codes),
            'gene': '; '.join(dict.fromkeys(g for g in genes if g)),
            'comment': comment,
        })

        tsv_data[f'{drug}_Status'] = 'Resistant'
        tsv_data[f'{drug}_Reported_Mutation'] = format_mutation(best)   # summary pick
        tsv_data[f'{drug}_All_Mutations'] = ', '.join(formatted)        # full list
        tsv_data[f'{drug}_Confidence'] = best_label
        tsv_data[f'{drug}_Consensus'] = (f"{vmap[best]['cons']:.3f}"
                                         if vmap[best]['cons'] is not None else 'NA')
    else:
        context[f'{drug}_detected'] = no_rt()
        context[f'{drug}_gene'] = 'None'
        context[f'{drug}_conf'] = 'N/A'
        tsv_data[f'{drug}_Status'] = 'Susceptible'
        tsv_data[f'{drug}_Reported_Mutation'] = 'None'
        tsv_data[f'{drug}_All_Mutations'] = 'None'
        tsv_data[f'{drug}_Confidence'] = 'N/A'
        tsv_data[f'{drug}_Consensus'] = 'NA'

resistant_drugs = [d for d in DRUGS if resistances.get(d)]
rset = set(resistant_drugs)

# ---- drug-resistance classification (WHO 2021 definitions) ----
#   MDR-TB     : isoniazid + rifampicin
#   RR-TB      : rifampicin
#   pre-XDR-TB : MDR/RR-TB + any fluoroquinolone
#   XDR-TB     : MDR/RR-TB + any fluoroquinolone + (bedaquiline or linezolid)

mdr  = {'Isoniazid', 'Rifampicin'} <= rset
rr   = 'Rifampicin' in rset
fq   = bool(FLUOROQUINOLONES & rset)
grpA = bool(NEW_SECONDLINE & rset)          # bedaquiline / linezolid

if (mdr or rr) and fq and grpA:
    cls_term, cls_full = 'XDR-TB', 'extensively drug-resistant tuberculosis'
elif (mdr or rr) and fq:
    cls_term, cls_full = 'pre-XDR-TB', 'pre-extensively drug-resistant tuberculosis'
elif mdr:
    cls_term, cls_full = 'MDR-TB', 'multidrug-resistant tuberculosis'
elif rr:
    cls_term, cls_full = 'RR-TB', 'rifampicin-resistant tuberculosis'
else:
    cls_term = cls_full = None

# ---- clinical interpretation as RichText: drug names + classification bolded ----
BODY = dict(font='Cambria', size=22)        # match report body (11 pt Cambria)
interp = RichText()
if resistant_drugs:
    names = [d.lower() for d in resistant_drugs]
    interp.add('Resistance-associated mutations were detected in genes associated with ', **BODY)
    for i, name in enumerate(names):
        if i:
            interp.add((' and ' if len(names) == 2 else ', and ') if i == len(names) - 1
                       else ', ', **BODY)
        interp.add(name, bold=True, **BODY)
    interp.add(' resistance. No resistance-associated mutations were detected for the '
               'remaining drugs tested.', **BODY)
    if cls_term:
        interp.add(f' Findings are consistent with {cls_full} (', **BODY)
        interp.add(cls_term, bold=True, **BODY)
        interp.add(').', **BODY)
    elif rset & FIRST_LINE:
        interp.add(' Findings indicate first-line resistance.', **BODY)
else:
    interp.add('No resistance-associated mutations were detected in the targeted regions. '
               'Findings are consistent with a pan-susceptible profile for the drugs tested.', **BODY)
context['clinical_interpretation'] = interp

context.update({
    'seq_platform':     os.environ.get('TB_PLATFORM', 'Oxford Nanopore'),
    'reference_genome': os.environ.get('TB_REFERENCE', 'H37Rv (NC_000962.3)'),
    'pipeline_version': os.environ.get('TB_PIPELINE_VER', 'v.1.1.0'),
    'db_version':       os.environ.get('TB_DB_VER', 'WHO catalogue v2'),
})


# --------------------------------------------------------------------------- #
#  Detailed Findings table
# --------------------------------------------------------------------------- #
def compute_qc(bam_path, bed_path, min_depth=30.0):
    """Per-amplicon mean depth + breadth of coverage from the BAM over BED regions.

    Returns (qc_summary_dict, amplicon_list) or (None, []) if it can't run.
    PASS for an amplicon = mean depth >= min_depth.
    """
    if not (bam_path and os.path.exists(bam_path) and os.path.exists(bed_path)):
        print(f"QC skipped (bam={bam_path}, bed={bed_path} missing)")
        return None, []

    # parse BED (whitespace-robust: file mixes tabs and spaces)
    regions = []
    with open(bed_path) as fh:
        for line in fh:
            line = line.strip()
            if not line or line.startswith(('#', 'track', 'browser')):
                continue
            parts = line.split()
            if len(parts) < 3:
                continue
            chrom, start, end = parts[0], int(parts[1]), int(parts[2])
            name = parts[3] if len(parts) > 3 else f"{chrom}:{start}-{end}"
            regions.append((chrom, start, end, name))

    try:
        bam = pysam.AlignmentFile(bam_path, "rb")
    except Exception as e:
        print(f"QC skipped (cannot open BAM): {e}")
        return None, []

    refs = set(bam.references)
    amplicons = []
    tot_bases = tot_depth = tot_ge = 0
    for chrom, start, end, name in regions:
        use_chrom = chrom if chrom in refs else (bam.references[0] if bam.references else None)
        length = end - start
        if use_chrom is None or length <= 0:
            continue
        cov = bam.count_coverage(use_chrom, start, end, quality_threshold=0)
        depths = [cov[0][i] + cov[1][i] + cov[2][i] + cov[3][i] for i in range(length)]
        mean_d = sum(depths) / length
        breadth = 100.0 * sum(1 for x in depths if x >= 1) / length
        ge = sum(1 for x in depths if x >= min_depth)
        amplicons.append({
            'name': name,
            'mean_depth': mean_d,
            'breadth': breadth,
            'status': 'PASS' if mean_d >= min_depth else 'FAIL',
        })
        tot_bases += length
        tot_depth += sum(depths)
        tot_ge += ge
    bam.close()

    if not amplicons:
        return None, []

    overall_mean = tot_depth / tot_bases if tot_bases else 0.0
    overall_cov = 100.0 * tot_ge / tot_bases if tot_bases else 0.0
    passed = [a['name'] for a in amplicons if a['status'] == 'PASS']
    failed = [a['name'] for a in amplicons if a['status'] == 'FAIL']
    summary = {
        'qc_status': 'PASS' if overall_mean >= min_depth else 'FAIL',
        'mean_depth': f"{overall_mean:.1f}x",
        'target_coverage': f"{overall_cov:.1f}% of target genes \u2265{min_depth:.0f}x",
        'amplicons_pass': f"{len(passed)}/{len(amplicons)}",
        'amplicons_fail': ', '.join(failed) if failed else 'None',
    }
    return summary, amplicons


def build_amplicon_table(doc, amplicons):
    """Fill the Amplicon Coverage table (clone donor row per amplicon)."""
    target = None
    for t in doc.tables:
        header = [c.text.strip().lower() for c in t.rows[0].cells]
        if any(h == 'amplicon' for h in header):
            target = t
            break
    if target is None:
        return
    donor_tr = target.rows[1]._tr
    parent = donor_tr.getparent()

    def fill(tr, values, status=None):
        for idx, (cell, val) in enumerate(zip(_Row(tr, target).cells, values)):
            p = cell.paragraphs[0]
            run = p.runs[0] if p.runs else p.add_run('')
            run.text = val
            for r in p.runs[1:]:
                r._element.getparent().remove(r._element)
            if idx == 3 and status:           # colour the QC column
                run.font.bold = True
                run.font.color.rgb = (RGBColor(0xC0, 0x00, 0x00) if status == 'FAIL'
                                      else RGBColor(0x2E, 0x7D, 0x32))

    if amplicons:
        for a in amplicons:
            new_tr = deepcopy(donor_tr)
            parent.append(new_tr)
            fill(new_tr,
                 [a['name'], f"{a['mean_depth']:.0f}", f"{a['breadth']:.1f}", a['status']],
                 status=a['status'])
    else:
        new_tr = deepcopy(donor_tr)
        parent.append(new_tr)
        fill(new_tr, ['Coverage data not available', '', '', ''])
    parent.remove(donor_tr)


def build_findings_table(doc, findings):
    """Clone the styled donor row (row 1) once per finding, then drop the donor."""
    target = None
    for t in doc.tables:
        header = [c.text.strip().lower() for c in t.rows[0].cells]
        if any('detected mutation' in h for h in header):
            target = t
            break
    if target is None:
        print("WARNING: Detailed Findings table not found in template.", file=sys.stderr)
        return

    donor_tr = target.rows[1]._tr
    parent = donor_tr.getparent()

    def fill(tr, values):
        for cell, val in zip(_Row(tr, target).cells, values):
            p = cell.paragraphs[0]
            if p.runs:
                p.runs[0].text = val
                for r in p.runs[1:]:
                    r._element.getparent().remove(r._element)
            else:
                p.add_run(val)

    if findings:
        for f in findings:
            new_tr = deepcopy(donor_tr)
            parent.append(new_tr)
            fill(new_tr, [f['drug'], f['mut'], f['gene'], f['comment']])
    else:
        new_tr = deepcopy(donor_tr)
        parent.append(new_tr)
        fill(new_tr, ['No resistance-associated mutations detected', '', '', ''])

    parent.remove(donor_tr)   # remove the sentinel donor row


# --------------------------------------------------------------------------- #
#  Render + save
# --------------------------------------------------------------------------- #
qc_summary, amplicon_qc = compute_qc(bam_arg, regions_bed, QC_MIN_DEPTH)
if qc_summary:
    print(f"QC: {qc_summary['qc_status']} | mean depth {qc_summary['mean_depth']} | "
          f"amplicons {qc_summary['amplicons_pass']} pass | failing: {qc_summary['amplicons_fail']}")
    context.update(qc_summary)
    tsv_data.update({
        'QC_Status': qc_summary['qc_status'],
        'Mean_Depth': qc_summary['mean_depth'],
        'Amplicons_Pass': qc_summary['amplicons_pass'],
        'Amplicons_Fail': qc_summary['amplicons_fail'],
    })
else:                                            # fall back to env / placeholders
    context.update({
        'qc_status':       os.environ.get('TB_QC_STATUS', 'Not reported'),
        'mean_depth':      os.environ.get('TB_MEAN_DEPTH', 'Not reported'),
        'target_coverage': os.environ.get('TB_TARGET_COV', 'Not reported'),
        'amplicons_pass':  os.environ.get('TB_AMPLICONS', 'Not reported'),
        'amplicons_fail':  os.environ.get('TB_AMPLICONS_FAIL', 'Not reported'),
    })

tpl = DocxTemplate(template_path)
tpl.render(context)                          # scalar + RichText tags
build_findings_table(tpl.docx, findings)     # variable tables, in Python
#build_amplicon_table(tpl.docx, amplicon_qc)

out_path = os.path.join(patient_dir, f'{output_base_name}_report.docx')
tpl.save(out_path)
print(f"Saved DOCX file to: {out_path}")

pd.DataFrame([tsv_data]).to_csv(
    os.path.join(patient_dir, f'{output_base_name}_results.tsv'), sep='\t', index=False)

print(f"\nResistance Profile Summary for {output_base_name}:")
print(f"Lineage: {tsv_data['Lineage']}")
for drug in DRUGS:
    s = tsv_data[f'{drug}_Status']
    if s == 'Resistant':
        rep = tsv_data[f'{drug}_Reported_Mutation']
        conf = tsv_data[f'{drug}_Confidence']
        cons = tsv_data[f'{drug}_Consensus']
        allm = tsv_data[f'{drug}_All_Mutations']
        print(f"  {drug}: Resistant (reported: {rep}, {conf}, consensus={cons}; all: {allm})")
    else:
        print(f"  {drug}: Susceptible")
